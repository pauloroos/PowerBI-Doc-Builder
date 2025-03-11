import shutil
import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image, ImageTk
import webbrowser
import os
import requests
import json
import ttkbootstrap as ttk  # Suporte a temas avançados
import tempfile
import subprocess
import time
import psutil
import pandas as pd
import clr
from docx import Document
from docx.shared import Inches  # Import necessário para ajustar o tamanho da imagem
from docx.shared import Pt, Cm
from docx.shared import RGBColor
from docx.oxml import parse_xml
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import nsdecls
from datetime import datetime
from pathlib import Path
import re
import google.generativeai as genai
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from core.helpers import *
from core.ai_description import *
from core.diagram_generator import *

def processar_pbix(pasta_pbix):
    # Cole aqui toda a lógica da função processar_pbix exatamente como está
    pass
# Função para processar arquivos PBIX
def processar_pbix(pasta_pbix):
    # Path Power BI Files
    pbix_folder = pasta_pbix
    
    # Carregar arquivo config e Atribuindo às variáveis
    config = carregar_config()
    cmd = config.get("cmd", "")
    ssas_dll = config.get("ssas_dll", "")
    pbi_desktop = config.get("pbi_desktop", "")

    # Create the folder "results" if it does not exist
    resultado_dir = os.path.join(pasta_pbix, "Resultado")
    os.makedirs(resultado_dir, exist_ok=True)  # Cria a pasta "Resultado" se não existir

    # DAX Query 
    dax_query = """
    DEFINE

    // Source Tables
    VAR __tables = INFO.TABLES()
    VAR __columns = INFO.COLUMNS()
    VAR __measures = INFO.MEASURES()
    VAR __formats = INFO.FORMATSTRINGDEFINITIONS()
    VAR __relationships = INFO.RELATIONSHIPS()
    VAR __calculationGroups = INFO.CALCULATIONGROUPS()
    VAR __calculationItems = INFO.CALCULATIONITEMS()
    VAR __dependencies = INFO.CALCDEPENDENCY()
    VAR __partitions = INFO.PARTITIONS()
    VAR __parametersquery = INFO.EXPRESSIONS()


    // Measures
    VAR __measuresPre = 
        ADDCOLUMNS(
                __measures,
                "Table",
            VAR TableID = [TableID] RETURN
                MAXX(FILTER(__tables, [ID] = TableID), [Name]),
                "Format",
            VAR MeasureID = [ID] RETURN
                MAXX(FILTER(__formats, [ObjectID] = MeasureID), [Expression])
        )

    VAR __measuresResult =
    SELECTCOLUMNS(
        __measuresPre,
        "tableName", [Table],
        "name", [Name],
        "expression", [Expression],
        "format", [FormatString],
        "isHidden", [IsHidden],
        "description", [Description],
        "type",
        SWITCH(
            [DataType],
            2, "String",
            6, "Whole Number",
            8, "Double",
            9, "Datetime",
            10, "Currency",
            11, "Boolean"
        ),
        "displayFolder", [DisplayFolder]
    ) 

    // Columns
    VAR __columnsPre =
    ADDCOLUMNS(
        FILTER(
            __columns,
            NOT CONTAINSSTRING([ExplicitName], "RowNumber")
        ),
        "TableName",
        VAR TableID = [TableID] RETURN
        MAXX(FILTER(__tables, [ID] = TableID), [Name]),
        "TableDescription",
        VAR TableID2 = [TableID] RETURN
        MAXX(FILTER(__tables, [ID] = TableID2), [Description])
    )

    VAR __columnsResult =
    SELECTCOLUMNS(
        __columnsPre,
        "tableName", [TableName],
        "tableDescription", [TableDescription],
        "name", COALESCE([ExplicitName], [InferredName]),
        "column", 
            "'" & [TableName] & "'" & 
            "[" & COALESCE([ExplicitName], [InferredName]) & "]",
        "sortedBy",
            VAR OrderID = [SortByColumnID] RETURN
            MAXX(
                FILTER(__columnsPre, [ID] = OrderID), 
                "'" & [TableName] & "'" & 
            "[" & COALESCE([ExplicitName], [InferredName]) & "]"),
        "format", COALESCE([FormatString], "String"),
        "displayFolder", [DisplayFolder],
        "isHidden", [IsHidden],
        "expression", [Expression]
    )

    // Relationships
    VAR relationshipsPre =
    ADDCOLUMNS(
        __relationships,
        "FromTable",
            VAR FromTable = [FromTableID] RETURN
            MAXX(FILTER(__columnsPre, [TableID] = FromTable), [TableName]),
        "ToTable",
            VAR ToTable = [ToTableID] RETURN
            MAXX(FILTER(__columnsPre, [TableID] = ToTable), [TableName]),
        "FromColumn",
            VAR FromColumn = [FromColumnID] RETURN
            MAXX(
                FILTER( __columnsPre, [ID] = FromColumn ),
                COALESCE( [ExplicitName], [InferredName] )),
        "ToColumn",
        VAR ToColumn = [ToColumnID]
        RETURN
            MAXX(
            FILTER( __columnsPre, [ID] = ToColumn ),
            COALESCE( [ExplicitName], [InferredName] )
            )
    )

    VAR __relationshipsResult =
    SELECTCOLUMNS(
        relationshipsPre,
        "from", "'" & [FromTable] & "'" & "[" & [FromColumn] & "]",
        "fromCardinality", IF([FromCardinality] = 2, "*", [FromCardinality]),
        "to",	"'" & [ToTable] & "'" & "[" & [ToColumn] & "]",
        "toCardinality", IF( [ToCardinality] = 2, "*", [ToCardinality]),
        "isActive", [IsActive],
        "isBidirectional", IF([CrossFilteringBehavior] = 2, "True", "False"),
        "relationship",
            VAR fromConcat = "'" & [FromTable] & "'" & "[" & [FromColumn] & "]"
            VAR toConcat = "'" & [ToTable] & "'" & "[" & [ToColumn] & "]"
            VAR fromCardinality = IF([FromCardinality] = 2, "✱", "1️⃣")
            VAR toCardinality = IF([ToCardinality] = 2, "✱", "1️⃣")
            VAR arrow = IF([CrossFilteringBehavior] = 2, "◀--▶", "◀--")
            RETURN
                fromConcat & " ( " & fromCardinality & " " & arrow & " " & toCardinality & " ) " & toConcat
    )

    // Calculation Groups
    VAR __calculationItemsPre =
    ADDCOLUMNS(
        ADDCOLUMNS(
        __calculationItems,
        "TableID",
            VAR CalculationGroup_ID = [CalculationGroupID] RETURN
            MAXX(FILTER(__calculationGroups, [ID] = CalculationGroup_ID), [TableID]),
        "Precedence",
            VAR CalculationGroup_ID = [CalculationGroupID] RETURN
            MAXX(FILTER(__calculationGroups, [ID] = CalculationGroup_ID), [Precedence])
        ),
        "CalculationGroup",
            VAR Table_ID = [TableID] RETURN
            MAXX(FILTER(__tables, [ID] = Table_ID), [Name]),
        "Format",
            VAR Format_ID = [FormatStringDefinitionID] RETURN
            MAXX(FILTER(__formats, [ID] = Format_ID), [Expression]),
        "CalculationGroupColumn",
            VAR Table_ID = [TableID]
            VAR Coluna_ID = MINX(FILTER(__columns, [TableID] = Table_ID), [ID])
            RETURN
                MAXX(
                    FILTER( __columns, [ID] = Coluna_ID ),
                    COALESCE( [ExplicitName], [InferredName] )
                ),
        "OrdinalColumn",
            VAR Table_ID = [TableID]
            VAR Coluna_ID = MAXX(FILTER(__columns, [TableID] = Table_ID), [ID])
            RETURN
                MAXX(
                    FILTER( __columns, [ID] = Coluna_ID ),
                    COALESCE( [ExplicitName], [InferredName] )
                )
    )

    VAR __calculationGroupsResult =
    SELECTCOLUMNS(
        __calculationItemsPre,
        "calculationGroup", [CalculationGroup],
        "precedence", [Precedence],
        "calculationItem", [Name],
        "expression", [Expression],
        "format", [Format],
        "ordinal", [Ordinal],
        "calculationGroupColumn", [CalculationGroupColumn],
        "ordinalColumn", [OrdinalColumn]
    )

    VAR __parametersResult =
    SELECTCOLUMNS(
        SUMMARIZE(
            FILTER(
                __dependencies,
                [REFERENCED_OBJECT_TYPE] = "M_EXPRESSION" && 
                CONTAINSSTRING( [REFERENCED_EXPRESSION], "IsParameterQuery" )
            ),
            [REFERENCED_OBJECT],
            [REFERENCED_EXPRESSION]
        ),
        "name", [REFERENCED_OBJECT],
        "expression", [REFERENCED_EXPRESSION],
        "description", MAXX(FILTER(__parametersquery, [Name] = [REFERENCED_OBJECT]), [Description])
    )


    // Queries M or DAX
    VAR __partitionsResult =
    SELECTCOLUMNS(
        __partitions,
        "tableName", 
        IF(
            NOT(ISBLANK([Name])) && LEN([Name]) > 37,
            LEFT([Name], LEN([Name]) - 37),
            [Name]
        ),
        "description", [Description],
        "queryDefinition", [QueryDefinition],
        "modifiedTime", [ModifiedTime],
        "refreshedTime", [RefreshedTime],
        "type", SWITCH( [Type], 4, "M", 2, "DAX", 7, "Internal" )
    )

    // Output
    EVALUATE __measuresResult ORDER BY [tableName] ASC, [name] ASC
    EVALUATE __columnsResult ORDER BY [tableName] ASC, [name] ASC
    EVALUATE __relationshipsResult ORDER BY [relationship] ASC
    EVALUATE __calculationGroupsResult ORDER BY [precedence] ASC, [ordinal] ASC
    EVALUATE __parametersResult 
    EVALUATE __partitionsResult ORDER BY [tableName] ASC 
    """

    def get_info_pro_datasets(cmd, dax_query, ssas_dll, pbi_desktop):
        
        # Load SSAS assembly
        clr.AddReference(ssas_dll)
        from Microsoft.AnalysisServices import Server  # type: ignore

        query = dax_query

        # Wait for Power BI Open
        def wait_for_powerbi_to_open(timeout=60, check_interval=5):
            """
            Waits for Power BI Desktop to open and start the Analysis Services server.

            timeout: Maximum wait time in seconds.
            check_interval: Interval between checks in seconds.
            """
            elapsed_time = 0
            while elapsed_time < timeout:
                server_name = get_powerbi_port()
                if server_name != "No instance of Power BI Desktop found.":
                    print(f"Power BI Desktop started! Server found at: {server_name}")
                    return server_name
                print(f"Waiting for Power BI Desktop to open... ({elapsed_time}s)")
                time.sleep(check_interval)
                elapsed_time += check_interval

            print("Timeout reached. Power BI Desktop did not start correctly.")
            return None

        def get_powerbi_port():
            # Search the process msmdsrv.exe (Analysis Services Power BI)
            for proc in psutil.process_iter(attrs=['pid', 'name']):
                if proc.info['name'] == "msmdsrv.exe":
                    pid = proc.info['pid']
                    
                    # Run netstat to capture active conections
                    result = subprocess.run(["netstat", "-ano"], capture_output=True, text=True)
                    connections = result.stdout.splitlines()
                    
                    # Filter corresponding line to PID of msmdsrv.exe
                    for line in connections:
                        if str(pid) in line and "LISTENING" in line:
                            match = re.search(r':(\d+)', line)
                            if match:
                                return f"localhost:{match.group(1)}"
            
            return "No instance of Power BI Desktop found."

        # Function to get the database ID
        def get_ssas_database_id(server_name):
            
            server = Server()
            try:
                server.Connect(f"Data Source={server_name}")
                # Assuming there is only one active database
                if server.Databases.Count > 0:
                    db_id = server.Databases[0].ID
                    return db_id
                else:
                    return None
            finally:
                server.Disconnect()

        def wait_for_powerbi_to_close(timeout=30, check_interval=2):
            """
            Waits for Power BI Desktop to close completely.
            """
            elapsed_time = 0
            while elapsed_time < timeout:
                if "PBIDesktop.exe" not in (p.name() for p in psutil.process_iter()):
                    print("Power BI Desktop closed successfully.")
                    return True
                print(f"Waiting for Power BI Desktop to close... ({elapsed_time}s)")
                time.sleep(check_interval)
                elapsed_time += check_interval
            print("Timeout waiting for Power BI Desktop to close.")
            return False

        # Iterate through each workspace
        for pbix_file_name in pbix_files_names:
            print(f"Processing: {pbix_file_name}.pbix")
            pbix_path = os.path.abspath(f"{pbix_folder}/{pbix_file_name}.pbix")

            # Execute
            if not os.path.exists(pbix_path):
                print(f"Error: The file {pbix_file_name}.pbix was not found.")
            else:
                try:
                    subprocess.Popen([pbi_desktop, pbix_path])
                    print(f"Power BI Desktop opened the file {pbix_file_name}.pbix.")
                except Exception as e:
                    print(f"Error opening Power BI Desktop: {e}")

            # Get Power BI Desktop Server
            server_name = wait_for_powerbi_to_open()

            if not server_name:
                print("Failed to start Power BI. Retrying...")
                time.sleep(5)
                server_name = wait_for_powerbi_to_open()

            if server_name:
                print(f"Server: {server_name}")
                database_id = get_ssas_database_id(server_name=server_name)
                if database_id:
                    print(f"Database: {database_id}")
                else:
                    print("No database found. Retrying after waiting...")
                    time.sleep(5)
                    database_id = get_ssas_database_id(server_name=server_name)
                    if not database_id:
                        print("Error: Unable to retrieve database ID.")
                        continue  # Skip this, try next
            else:
                print("No server found. Skipping this report.")
                continue

            # Dax Studio Variables
            server = server_name
            database = database_id

            # Output path
            output_dir = os.path.join(resultado_dir, "Arquivos", pbix_file_name)
            os.makedirs(output_dir, exist_ok=True)  # Garante que a pasta do dataset será criada
            output_path = f"{output_dir}/extract.csv"

            # Clean up any existing extract files before execution
            for old_name in ["extract.csv", "extract_2.csv", "extract_3.csv", 
                            "extract_4.csv", "extract_5.csv", "extract_6.csv"]:
                old_path = os.path.join(output_dir, old_name)
                if os.path.exists(old_path):
                    os.remove(old_path)
                    print(f"    Removed existing source file: {old_name}")

            max_retries = 3
            retry_delay = 5

            for attempt in range(max_retries):
                try:
                    # Executing
                    subprocess.run([
                        cmd, "csv", output_path,
                        "-s", server,
                        "-d", database,
                        "-q", query
                    ], check=True)
                    break
                except subprocess.CalledProcessError:
                    print(f"Error executing dscmd. Attempt {attempt + 1} of {max_retries}")
                    if attempt < max_retries - 1:
                        time.sleep(retry_delay)
                    else:
                        print("Persistent failure executing dscmd. Skipping this report.")


            # Close Power BI Desktop
            subprocess.run(["taskkill", "/F", "/IM", "PBIDesktop.exe"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            print("Power BI Desktop was closed.")
            
            
            # Remaping the extracted files
            rename_map = {
                "extract.csv": "measures.csv",
                "extract_2.csv": "columns.csv",
                "extract_3.csv": "relationships.csv",
                "extract_4.csv": "calculation_groups.csv",
                "extract_5.csv": "parameters.csv",
                "extract_6.csv": "partitions.csv"
            }

            # Rename files
            for old_name, new_name in rename_map.items():
                old_path = os.path.join(output_dir, old_name)
                new_path = os.path.join(output_dir, new_name)

                # Check if old file exists
                if os.path.exists(old_path):
                    # If new file already exists, remove it first
                    if os.path.exists(new_path):
                        os.remove(new_path)
                        print(f"    Removed existing file: {new_name}")
                
                    os.rename(old_path, new_path)
                    print(f"    Renamed: {old_name} -> {new_name}")
                else:
                    print(f"    File not found: {old_name}")

            
            # Wait for the process to actually terminate before opening the next file
            time.sleep(5)
            wait_for_powerbi_to_close()


        print("Processing completed all pbix files in local folder.")

    def create_documentation():
        def create_semantic_model_doc(dataset_name):
            output_dir = os.path.join(resultado_dir, "Arquivos", pbix_file_name)
            directory = os.path.join(resultado_dir, "Arquivos", dataset_name)
            doc = Document()
            
            # Criar cabeçalho
            section = doc.sections[0]
            header = section.header
            header_table = header.add_table(rows=1, cols=2, width=Inches(6))
            hdr_cells = header_table.rows[0].cells
            
            # Baixar a logo temporariamente
            # logo_path = download_logo(logo_url)
            # logo_path = os.path.join(get_base_path(), "assets", "seu-logo.png")
            logo_path = os.path.join(get_base_path(), "assets", "seu-logo.png")
            
            # Inserir a logo na célula da esquerda
            if os.path.exists(logo_path):  # Verifica se a logo existe antes de inserir
                # Limpar o conteúdo da célula antes de adicionar a imagem
                hdr_cells[0]._element.clear_content()
                
                # Adicionar a imagem na célula
                paragraph = hdr_cells[0].add_paragraph()
                run = paragraph.add_run()
                run.add_picture(logo_path, width=Inches(1))  # Ajuste o tamanho conforme necessário
                

            # Inserir o título na célula da direita
            hdr_cells[1].paragraphs[0].add_run("Documentação Técnica do Modelo Semântico").bold = True
            hdr_cells[1].paragraphs[0].alignment = 2  # Alinhar à direita (2 = RIGHT)

            # Ajustar a largura das células para evitar sobreposição
            hdr_cells[0].width = Inches(2)  # Largura da célula da logo
            hdr_cells[1].width = Inches(6)  # Largura da célula do título
            
            
            # Ajustar as margens do documento
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(0)
                section.bottom_margin = Cm(0)
                section.left_margin = Cm(1)
                section.right_margin = Cm(1)
                
                # Criar rodapé
                footer = section.footer
                footer_paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                
                # Adicionar a data de geração do documento
                now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                footer_paragraph.text = f"Documento gerado em: {now}"
                footer_paragraph.alignment = 2
                footer_paragraph.style.font.size = Pt(6)

            # Custom Style
            style = doc.styles['Normal']
            style.font.name = 'Segoe UI'
            style.font.size = Pt(10)
            style.paragraph_format.space_after = Pt(2)
            
            # Ajustar tamanho da fonte dos títulos manualmente
            def set_heading_font_size(paragraph, font_size):
                """Modifica o tamanho da fonte de um título."""
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    
            # Main Title
            # doc.add_heading("Semantic Model Technical Documentation", level=1)
            # titulo = doc.add_heading("Documentação Técnica do Modelo Semantico", level=1)
            # set_heading_font_size(titulo, 16)  # Definir fonte 16 para o título principal
            
            # Criar a capa
            titulo_capa = doc.add_paragraph()
            titulo_capa.alignment = 1  # Centralizado
            titulo_run = titulo_capa.add_run("\n\n\n\n\n\nDocumentação Técnica do Modelo Semântico\n\n")
            titulo_run.bold = True
            titulo_run.font.size = Pt(22)

            # Buscar o nome do modelo para exibir abaixo do título
            modelo_nome = dataset_name  # O nome do modelo é passado como parâmetro

            modelo_run = titulo_capa.add_run(f"{modelo_nome}")
            modelo_run.font.size = Pt(16)

            # Adicionar uma quebra de página após a capa
            doc.add_page_break()
            
            # Criar a seção "Sobre esta Documentação"
            doc.add_heading("Sobre esta Documentação", level=1)

            doc.add_paragraph(
                "Este documento tem como objetivo fornecer uma visão detalhada do modelo semântico "
                "extraído do Power BI. A seguir, cada seção deste documento é explicada:"
            )

            doc.add_heading("1. Overview", level=2)
            doc.add_paragraph(
                "A seção Overview fornece informações gerais sobre o modelo semântico, incluindo "
                "seu nome e uma breve descrição sobre sua finalidade. Caso esteja configurado corretamente, uma descrição via IA também será incluída."
            )

            doc.add_heading("2. Diagrama", level=2)
            doc.add_paragraph(
                "A seção Diagrama apresenta uma visualização gráfica do modelo semântico no formato de diagrama relacional. "
                "Este diagrama destaca as conexões entre as tabelas, incluindo as cardinalidades e direções dos relacionamentos. "
                "Ele é especialmente útil para compreender a estrutura lógica do modelo e o fluxo de dados entre as tabelas."
            )

            doc.add_heading("3. Colunas", level=2)
            doc.add_paragraph(
                "Nesta seção, cada tabela do modelo é apresentada com seu nome e, quando disponível, sua descrição. "
                "Para cada tabela, são listadas as colunas associadas, incluindo o nome da coluna, a coluna utilizada para ordenação (caso aplicável), o formato de exibição e, "
                "informações sobre tipo de dado, formato e se a coluna está oculta ou não."
            )

            doc.add_heading("4. Medidas", level=2)
            doc.add_paragraph(
                "A seção de Medidas contém todas as medidas DAX criadas no modelo. Cada medida inclui "
                "sua expressão, formato e descrição, se disponível."
            )

            doc.add_heading("5. Relacionamentos", level=2)
            doc.add_paragraph(
                "Lista todos os relacionamentos entre as tabelas do modelo, indicando quais colunas "
                "estão relacionadas, a cardinalidade da relação e se ela está ativa ou não."
            )

            doc.add_heading("6. Partições", level=2)
            doc.add_paragraph(
                "Esta seção detalha como cada tabela do modelo foi construída, listando o nome da tabela, a linguagem utilizada (M ou DAX) "
                "e a expressão completa usada para sua criação. Essa informação é essencial para entender a origem dos dados e o método de carregamento adotado no modelo."
            )

            doc.add_heading("7. Parâmetros", level=2)
            doc.add_paragraph(
                "Lista os parâmetros do modelo, caso existam. Cada parâmetro é acompanhado de sua "
                "expressão DAX ou M e sua finalidade."
            )
            
            # Adicionar uma quebra de página após a explicação
            doc.add_page_break()         
            
            # Overview
            subtitulo = doc.add_heading("Overview", level=2)
            set_heading_font_size(subtitulo, 14)  # Definir fonte 14
            doc.add_paragraph(f"Modelo Semantico: {dataset_name}")
            doc.add_paragraph("Este documento descreve a estrutura do modelo semântico extraído, incluindo tabelas, colunas, relacionamentos, partições, medidas, grupos de cálculo e parâmetros.")
            
            # 🔹 Gerar a descrição via IA usando os CSVs já extraídos
            descricao_ia = gerar_descricao_ia(directory)
            
            # Ignorar mensagens de erro visíveis no docx
            if descricao_ia.startswith("❌"):
                descricao_ia = "Descrição não disponível. Verifique a chave da API Gemini em sua configuração."

            # 🔹 Adicionar ao documento (como parágrafo separado)
            doc.add_paragraph("Descrição automática do modelo (gerada com IA):", style="Intense Quote")
            add_texto_formatado(doc, descricao_ia)
            # doc.add_paragraph(descricao_ia)            
            
            # 🔹 Adiciona uma nova página para o diagrama
            doc.add_page_break()
            doc.add_heading("Diagrama", level=1)

            # 🔹 Gera a imagem do diagrama
            caminho_csv = os.path.join(directory, "relationships.csv")
            pasta_documentacao = os.path.join(resultado_dir, "Documentacao")
            imagem_diagrama = gerar_diagrama(
                path_csv=caminho_csv,
                nome_base=dataset_name,
                pasta_documentacao=pasta_documentacao
            )

            # 🔹 Insere a imagem no documento com redimensionamento proporcional (máx 20x22 cm)
            if os.path.exists(imagem_diagrama):
                img = Image.open(imagem_diagrama)
                w_px, h_px = img.size
                dpi = 96  # padrão comum
                w_cm = w_px / dpi * 2.54
                h_cm = h_px / dpi * 2.54

                max_w = 20
                max_h = 22
                escala = min(max_w / w_cm, max_h / h_cm)

                largura_final = w_cm * escala
                altura_final = h_cm * escala

                paragrafo = doc.add_paragraph()
                paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragrafo.add_run()
                run.add_picture(imagem_diagrama, width=Cm(largura_final), height=Cm(altura_final))
                        
            # Load csv files
            file_paths = {
                "Calculation Groups": os.path.join(directory, "calculation_groups.csv"),
                "Columns": os.path.join(directory, "columns.csv"),
                "Measures": os.path.join(directory, "measures.csv"),
                "Parameters": os.path.join(directory, "parameters.csv"),
                "Partitions": os.path.join(directory, "partitions.csv"),
                "Relationships": os.path.join(directory, "relationships.csv"),
            }
            
            dataframes = {name: pd.read_csv(path, encoding="utf-8", delimiter=";") for name, path in file_paths.items() if os.path.exists(path)}

            # Nan to ""
            for key in dataframes:
                dataframes[key] = dataframes[key].fillna("")

            # Apply style tables
            def apply_table_style(table):
                for i, row in enumerate(table.rows):
                    for cell in row.cells:
                        if cell.paragraphs:
                            if i == 0:
                                cell.paragraphs[0].runs[0].font.size = Pt(9)
                                cell.paragraphs[0].runs[0].font.bold = True
                                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(225, 225, 225)
                            else:
                                cell.paragraphs[0].runs[0].font.size = Pt(8)
                            
                        if i == 0:
                            cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="3E3E61"/>'.format(nsdecls('w'))))
                        elif i % 2 == 0:
                            cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="F3F3F3"/>'.format(nsdecls('w'))))

            # Add sections
            sections = {
                "Columns": ["Columns", ["Table", "Table Description", "Column", "Sorted By", "Format", "Expression"], ["tableName","tableDescription","name","sortedBy","format","expression"]],
                "Partitions": ["Partitions", ["Table", "Type", "Query"], ["tableName", "type", "queryDefinition"]],
                "Relationships": ["Relationships", ["Relationship", "isActive", "isBidirectional"], ["relationship", "isActive", "isBidirectional"]],
                "Measures": ["Measures", ["Name", "Expression", "Description", "Format"], ["name", "expression", "description", "format"]],
                "Calculation Groups": ["Calculation Groups", ["Group", "Item", "Expression", "Ordinal", "Format"], ["calculationGroup", "calculationItem", "expression", "ordinal", "format"]],
                "Parameters": ["Parameters", ["Name", "Description", "Expression"], ["name","description","expression"]],
            }
            
            for key, (title, headers, columns) in sections.items():
                if key not in dataframes or dataframes[key].empty:
                    continue
                
                if key in dataframes:      
                    doc.add_page_break()  # Adiciona uma quebra de página antes de cada seção
                    titulo_formatado = titulos_pt.get(title, title)  # usa o valor traduzido se existir
                    subtitulo = doc.add_heading(titulo_formatado, level=2)
                    set_heading_font_size(subtitulo, 14)  # Definir fonte 14
                    table = doc.add_table(rows=1, cols=len(headers))
                    table.autofit = False
                    # Set column widths
                    
                    if key == "Partitions":
                        col_widths = [Cm(6), Cm(2), Cm(12)]  # Ajuste conforme necessário
                        for i, column in enumerate(table.columns):
                            column.width = col_widths[i]
                        # table.columns[0].width = Cm(1)  # Table
                        # table.columns[1].width = Cm(1)  # Type
                        # table.columns[2].width = Cm(48)  # Query

                    elif key == "Relationships":
                        col_widths = [Cm(15), Cm(2), Cm(3)]  # Ajuste conforme necessário
                        for i, column in enumerate(table.columns):
                            column.width = col_widths[i]
                        # table.columns[0].width = Cm(24)  # Relationship
                        # table.columns[1].width = Cm(1)  # isActive
                        # table.columns[2].width = Cm(1)  # isBidirectional

                    elif key == "Measures":
                        col_widths = [Cm(3), Cm(11), Cm(3), Cm(3)]  # Ajuste conforme necessário
                        for i, column in enumerate(table.columns):
                            column.width = col_widths[i]
                        # table.columns[0].width = Cm(2)  # Name
                        # table.columns[1].width = Cm(36)  # Expression
                        # table.columns[2].width = Cm(36)  # Description
                        # table.columns[3].width = Cm(2)  # Format

                    elif key == "Calculation Groups":
                        col_widths = [Cm(2), Cm(2), Cm(13), Cm(1), Cm(2)]  # Ajuste conforme necessário
                        for i, column in enumerate(table.columns):
                            column.width = col_widths[i]
                        # table.columns[0].width = Cm(2)  # Group
                        # table.columns[1].width = Cm(2)  # Item
                        # table.columns[2].width = Cm(36)  # Expression
                        # table.columns[3].width = Cm(1)  # Ordinal
                        # table.columns[4].width = Cm(2)  # Format

                    elif key == "Parameters":
                        col_widths = [Cm(3), Cm(6), Cm(11)]  # Ajuste conforme necessário
                        for i, column in enumerate(table.columns):
                            column.width = col_widths[i]
                        # table.columns[0].width = Cm(1)  # Name
                        # table.columns[1].width = Cm(3)  # Description
                        # table.columns[2].width = Cm(36)  # Expression
                        
                    if key == "Columns":
                        # Agrupar as colunas por tabela
                        df_columns = dataframes[key]
                        tables = [
                            table for table in df_columns["tableName"].unique()
                            if not (table.startswith("DateTableTemplate") or table.startswith("LocalDateTable"))
                        ] # Obtém os nomes das tabelas únicas

                        for table_name in tables:
                            titulo_tabela = doc.add_heading(table_name, level=3)  # Subtítulo para cada tabela
                            set_heading_font_size(titulo_tabela, 12)  # Definir fonte 11
                            
                            # Obter a descrição da tabela antes de criar a tabela
                            table_description = df_columns[df_columns["tableName"] == table_name]["tableDescription"].iloc[0] if "tableDescription" in df_columns.columns else ""
                            
                            # Adiciona a linha de descrição antes da tabela
                            descricao_paragrafo = doc.add_paragraph()
                            descricao_paragrafo.add_run("Descrição da Tabela: ").bold = True # Adiciona negrito
                            descricao_paragrafo.add_run(table_description)  # Concatenar a descrição
                                
                            table = doc.add_table(rows=1, cols=len(headers) - 2)  # Criar tabela sem a coluna "tableName" e "tableDescription"
                            table.autofit = False
                            
                            # Ajustar largura das colunas de forma padronizada
                            col_widths = [Cm(6), Cm(3), Cm(3), Cm(8)]  # Ajuste conforme necessário
                            for i, column in enumerate(table.columns):
                                column.width = col_widths[i]
                            
                            # Definir cabeçalhos
                            hdr_cells = table.rows[0].cells
                            for i, header in enumerate(headers[2:]):  # Ignorando "Table" no cabeçalho
                                nome_formatado = headers_pt.get(header, header)  # usa tradução se existir
                                hdr_cells[i].text = nome_formatado
                                hdr_cells[i].width = col_widths[i]  # 🔹 Aqui aplica a largura da célula
                                hdr_cells[i]._tc.width = col_widths[i]

                            # Adicionar as colunas correspondentes a essa tabela
                            for _, row in df_columns[df_columns["tableName"] == table_name].iterrows():
                                row_cells = table.add_row().cells
                                for i, col in enumerate(columns[2:]):  # Ignorando "tableName" nos dados
                                    if i < len(row_cells):
                                        row_cells[i].text = str(row[col]) if col in row else ""
                                        row_cells[i].width = col_widths[i]  # 🔹 Também aplica nas células da linha
                                        row_cells[i]._tc.width = col_widths[i]

                            apply_table_style(table)  # Aplicar estilo na tabela

                    else:                      
                        # Adicionar cabeçalhos
                        hdr_cells = table.rows[0].cells
                        for i, header in enumerate(headers):
                            nome_formatado = headers_pt.get(header, header)  # usa tradução se existir
                            hdr_cells[i].text = nome_formatado
                            hdr_cells[i].width = col_widths[i]  # 🔹 Aqui aplica a largura da célula
                            hdr_cells[i]._tc.width = col_widths[i]

                        for _, row in dataframes[key].iterrows():
                            row_cells = table.add_row().cells
                            for i, col in enumerate(columns):
                                if i < len(row_cells):
                                    row_cells[i].text = str(row[col]) if col in row else ""
                                    row_cells[i].width = col_widths[i]  # 🔹 Também aplica nas células da linha
                                    row_cells[i]._tc.width = col_widths[i]

                        apply_table_style(table)       
            
            # Ensure the directory exists
            dir_path = os.path.join(resultado_dir, "Documentacao")
            os.makedirs(dir_path, exist_ok=True)  # Cria a pasta de documentação
            
            file_path = f"{dir_path}/{dataset_name}.docx"
            doc.save(file_path)
            print(f"File {file_path} saved.")

        for pbix_file_name in pbix_files_names:
            print(f"Processing report {pbix_file_name}")
            create_semantic_model_doc(dataset_name=pbix_file_name)

        print("All docx files were generated!")

    # Exec
    pbix_files_names = [os.path.splitext(f)[0] for f in os.listdir(pbix_folder) if f.endswith('.pbix')]  
    get_info_pro_datasets(cmd, dax_query, ssas_dll, pbi_desktop)
    create_documentation()    
