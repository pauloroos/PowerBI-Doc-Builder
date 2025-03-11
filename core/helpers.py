
import os
import json
import sys
import requests
import tempfile
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

CONFIG_FILE = "config.json"

# Variáveis para os links (edite aqui com seus dados)
LINKEDIN_URL = "https://www.linkedin.com/in/pauloroosf/"
GITHUB_URL = "https://github.com/pauloroos"
EMAIL_URL = "mailto:pauloroosf@hotmail.com"
SITE_URL = "https://pauloroosf.dev"

# Definir caminho da logo (Substitua pelo caminho correto da sua imagem)
logo_url  = "URL_LOGO"

# Caminhos padrão para os executáveis
VALORES_PADRAO = {
    "cmd": "C:\\Program Files\\DAX Studio\\dscmd.exe",
    "ssas_dll": "C:\\Program Files\\DAX Studio\\bin\\Microsoft.AnalysisServices.dll",
    "pbi_desktop": "C:\\Program Files\\WindowsApps\\Microsoft.MicrosoftPowerBIDesktop_2.140.1351.0_x64__8wekyb3d8bbwe\\bin\\PBIDesktop.exe",
    "api_key": "SUA_CHAVE_API_GEMINI"
}

# URLs dos ícones sociais
ICON_URLS = {
    "LinkedIn": "https://cdn-icons-png.flaticon.com/512/145/145807.png",
    "GitHub": "https://cdn-icons-png.flaticon.com/512/2111/2111432.png",  # Ícone branco do GitHub
    "E-mail": "https://cdn-icons-png.flaticon.com/512/552/552486.png",  # Ícone de email genérico
    "Site": "https://cdn-icons-png.flaticon.com/512/3178/3178285.png"  # Ícone branco para sites
}


# Traduções para o português
titulos_pt = {
    "Columns": "Colunas",
    "Partitions": "Partições",
    "Relationships": "Relacionamentos",
    "Measures": "Medidas",
    "Calculation Groups": "Grupos de Cálculo",
    "Parameters": "Parâmetros"
}
headers_pt = {
    "Table": "Tabela",
    "Table Description": "Descrição da Tabela",
    "Column": "Coluna",
    "Sorted By": "Ordenado por",
    "Format": "Formato",
    "Expression": "Expressão",
    "Display Folder": "Pasta de Exibição",
    "Hidden": "Oculta",
    "Name": "Nome",
    "Description": "Descrição",
    "Relationship": "Relacionamento",
    "isActive": "Está Ativo",
    "isBidirectional": "Bidirecional",
    "Type": "Tipo",
    "Query": "Consulta",
    "Group": "Grupo",
    "Item": "Item",
    "Ordinal": "Ordem",
    "calculationGroup": "Grupo de Cálculo",
    "calculationItem": "Item de Cálculo",
    "ordinal": "Ordem",
    "calculationGroupColumn": "Coluna de Grupo",
    "ordinalColumn": "Coluna de Ordem",
}

# Carregar configuração do JSON
def carregar_config():
    """Carrega o config.json. Se não existir, cria com valores padrão."""
    if not os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, "w", encoding="utf-8") as f:
            json.dump(VALORES_PADRAO, f, indent=4)
        return VALORES_PADRAO.copy()
    
    with open(CONFIG_FILE, "r", encoding="utf-8") as f:
        return json.load(f)

# Salvar configuração no JSON
def salvar_config(config):
    with open(CONFIG_FILE, "w", encoding="utf-8") as file:
        json.dump(config, file, indent=4)

# Formatar texto Gerador pela IA
def add_texto_formatado(doc, texto):
    texto = texto.replace("`", "'")  # Substituir crase por aspas simples

    linhas = texto.splitlines()
    for linha in linhas:
        linha = linha.strip()
        if not linha:
            continue

        if linha.startswith("* "):  # Bullet point
            item = doc.add_paragraph(style="List Bullet")
            conteudo = linha[2:]  # remove "* "
        else:
            item = doc.add_paragraph()
            conteudo = linha

        # Processar negrito (**texto**)
        i = 0
        while i < len(conteudo):
            if conteudo[i:i+2] == "**":
                i += 2
                negrito_texto = ""
                while i < len(conteudo) and conteudo[i:i+2] != "**":
                    negrito_texto += conteudo[i]
                    i += 1
                i += 2  # pular os ** finais
                run = item.add_run(negrito_texto)
                run.bold = True
            else:
                run = item.add_run(conteudo[i])
                i += 1

# Dowload do logo
def download_logo(url):
    response = requests.get(url, stream=True)
    if response.status_code == 200:
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        with open(temp_file.name, "wb") as f:
            for chunk in response.iter_content(1024):
                f.write(chunk)
        return temp_file.name
    return None

# Obter caminho base
def get_base_path():
    if getattr(sys, 'frozen', False):
        return sys._MEIPASS
    return os.path.dirname(__file__)

# Função para centralizar a janela
def centralizar_janela(janela, largura, altura):
    """Centraliza a janela no meio da tela."""
    janela.update_idletasks()  # Atualiza a geometria da janela
    w_screen = janela.winfo_screenwidth()
    h_screen = janela.winfo_screenheight()
    x = (w_screen - largura) // 2
    y = (h_screen - altura) // 2
    janela.geometry(f"{largura}x{altura}+{x}+{y}")  # Largura x Altura + X + Y


def set_heading_font_size(paragraph, font_size):
    for run in paragraph.runs:
        run.font.size = Pt(font_size)

def apply_table_style(table):
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if cell.paragraphs:
                if i == 0:
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(225, 225, 225)
                else:
                    cell.paragraphs[0].runs[0].font.size = Pt(8)
            if i == 0:
                cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="3E3E61"/>'.format(nsdecls('w'))))
            elif i % 2 == 0:
                cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="F3F3F3"/>'.format(nsdecls('w'))))
